#!/usr/bin/env python3
"""
genera_psc_batch.py
Genera 9 PSC DOCX personalizzati per i cantieri ITEA Via Doss Trento, Trento.
Uso: python genera_psc_batch.py

Dipendenze:
    pip install python-docx pymupdf requests
"""

import os, sys, re, shutil, json, textwrap, requests
from pathlib import Path

try:
    from docx import Document
except ImportError:
    sys.exit("Installare python-docx:  pip install python-docx")

try:
    import fitz  # PyMuPDF
except ImportError:
    sys.exit("Installare PyMuPDF:  pip install pymupdf")

# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURAZIONE
# ══════════════════════════════════════════════════════════════════════════════
ANTHROPIC_KEY = os.environ.get("ANTHROPIC_KEY", "")  # imposta la variabile d'ambiente ANTHROPIC_KEY
MODEL_CLAUDE  = "claude-haiku-4-5-20251001"

BASE          = Path(r"C:\Users\franz\OneDrive\Desktop\PSC ITEA per Frenz")
TEMPLATE_DOCX = BASE / "02 PSC e XLS di partenza" / "PSC_UI_43103914.docx"
DATI_DIR      = BASE / "03 Progetto e dati ITEA"  / "trento_doss trento"
OUTPUT_DIR    = Path(r"C:\Users\franz\safetyai\PSC_OUTPUT")

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ══════════════════════════════════════════════════════════════════════════════
# UTILITÀ FORMATO
# ══════════════════════════════════════════════════════════════════════════════
def piano_str(p: str) -> str:
    """'5' → '5°'"""
    return {"1":"1°","2":"2°","3":"3°","4":"4°","5":"5°","6":"6°"}.get(str(p), f"{p}°")

def fmt_euro(v: float) -> str:
    """17534.71 → '17.534,71'"""
    intero, dec = f"{v:.2f}".split(".")
    gruppi = []
    while len(intero) > 3:
        gruppi.insert(0, intero[-3:])
        intero = intero[:-3]
    gruppi.insert(0, intero)
    return ".".join(gruppi) + "," + dec

def gg_to_mesi(gg: int) -> str:
    """120 → '4 mesi';  60 → '2 mesi'"""
    m = round(gg / 30)
    return f"{m} mesi"

# ══════════════════════════════════════════════════════════════════════════════
# DATI DEI 9 CANTIERI
# ══════════════════════════════════════════════════════════════════════════════
CANTIERI = [
    {
        "codice_ui"      : "43103914",
        "cod_ed"         : "958",
        "civico"         : "45/1",
        "scala"          : "A",
        "interno"        : "23",
        "piano"          : "5",
        "p_ed"           : "5620",
        "c_c"            : "Trento",
        "sup"            : "99,81",
        "anno_costr"     : "1983",
        "amianto"        : False,
        "importo"        : 17534.71,
        "oneri"          : 3109.31,
        "oneri_lettere"  : "tremilacentonove/31",
        "durata_gg"      : 120,
        "data_inizio"    : "8 giugno 2026",
        "tipo_intervento": "sistemazione parziale",
    },
    {
        "codice_ui"      : "43104709",
        "cod_ed"         : "1346",
        "civico"         : "37",
        "scala"          : "A",
        "interno"        : "8",
        "piano"          : "2",
        "p_ed"           : "2496/3",
        "c_c"            : "Trento",
        "sup"            : "37,29",
        "anno_costr"     : "1988",
        "amianto"        : False,
        "importo"        : 32056.73,
        "oneri"          : 3500.00,
        "oneri_lettere"  : "tremilasettecento/00",
        "durata_gg"      : 120,
        "data_inizio"    : "22 giugno 2026",
        "tipo_intervento": "sistemazione completa",
    },
    {
        "codice_ui"      : "43104715",
        "cod_ed"         : "1346",
        "civico"         : "37",
        "scala"          : "A",
        "interno"        : "14",
        "piano"          : "3",
        "p_ed"           : "2496/3",
        "c_c"            : "Trento",
        "sup"            : "37,29",
        "anno_costr"     : "1988",
        "amianto"        : False,
        "importo"        : 31717.14,
        "oneri"          : 3500.00,
        "oneri_lettere"  : "tremilasettecento/00",
        "durata_gg"      : 120,
        "data_inizio"    : "6 luglio 2026",
        "tipo_intervento": "sistemazione completa",
    },
    {
        "codice_ui"      : "43105028",
        "cod_ed"         : "1421",
        "civico"         : "35",
        "scala"          : "A",
        "interno"        : "4",
        "piano"          : "1",
        "p_ed"           : "2496/6",
        "c_c"            : "Trento",
        "sup"            : "28,59",
        "anno_costr"     : "1989",
        "amianto"        : False,
        "importo"        : 9067.72,
        "oneri"          : 2000.00,
        "oneri_lettere"  : "duemila/00",
        "durata_gg"      : 60,
        "data_inizio"    : "20 luglio 2026",
        "tipo_intervento": "sistemazione parziale",
    },
    {
        "codice_ui"      : "43105031",
        "cod_ed"         : "1421",
        "civico"         : "35",
        "scala"          : "A",
        "interno"        : "7",
        "piano"          : "2",
        "p_ed"           : "2496/3",   # ← da verificare
        "c_c"            : "Trento",
        "sup"            : "78,55",
        "anno_costr"     : "1989",
        "amianto"        : False,
        "importo"        : 42198.01,
        "oneri"          : 4000.00,
        "oneri_lettere"  : "quattromila/00",
        "durata_gg"      : 140,
        "data_inizio"    : "3 agosto 2026",
        "tipo_intervento": "sistemazione completa",
    },
    {
        "codice_ui"      : "43105049",
        "cod_ed"         : "1421",
        "civico"         : "35",
        "scala"          : "A",
        "interno"        : "25",
        "piano"          : "5",
        "p_ed"           : "2496/6",
        "c_c"            : "Trento",
        "sup"            : "78,54",
        "anno_costr"     : "1989",
        "amianto"        : False,
        "importo"        : 46801.38,
        "oneri"          : 4200.00,
        "oneri_lettere"  : "quattromiladuecento/00",
        "durata_gg"      : 140,
        "data_inizio"    : "17 agosto 2026",
        "tipo_intervento": "sistemazione completa",
    },
    {
        "codice_ui"      : "43105448",
        "cod_ed"         : "1487",
        "civico"         : "31",
        "scala"          : "A",
        "interno"        : "5",
        "piano"          : "1",
        "p_ed"           : "6047",
        "c_c"            : "Trento",
        "sup"            : "78,38",
        "anno_costr"     : "1990",
        "amianto"        : True,   # bonifica amianto prevista
        "importo"        : 48277.21,
        "oneri"          : 5500.00,
        "oneri_lettere"  : "cinquemilacinquecento/00",
        "durata_gg"      : 160,
        "data_inizio"    : "31 agosto 2026",
        "tipo_intervento": "ristrutturazione parziale",
    },
    {
        "codice_ui"      : "43105459",
        "cod_ed"         : "1487",
        "civico"         : "31",
        "scala"          : "A",
        "interno"        : "16",
        "piano"          : "4",
        "p_ed"           : "6047",
        "c_c"            : "Trento",
        "sup"            : "77,31",
        "anno_costr"     : "1990",
        "amianto"        : False,
        "importo"        : 46068.67,
        "oneri"          : 4200.00,
        "oneri_lettere"  : "quattromiladuecento/00",
        "durata_gg"      : 140,
        "data_inizio"    : "14 settembre 2026",
        "tipo_intervento": "ristrutturazione parziale",
    },
    {
        "codice_ui"      : "43105460",
        "cod_ed"         : "1487",
        "civico"         : "31",
        "scala"          : "A",
        "interno"        : "17",
        "piano"          : "4",
        "p_ed"           : "6047",
        "c_c"            : "Trento",
        "sup"            : "46,23",
        "anno_costr"     : "1990",
        "amianto"        : False,
        "importo"        : 33434.18,
        "oneri"          : 3600.00,
        "oneri_lettere"  : "tremilaseicento/00",
        "durata_gg"      : 100,
        "data_inizio"    : "28 settembre 2026",
        "tipo_intervento": "sistemazione completa",
    },
]

# ══════════════════════════════════════════════════════════════════════════════
# LETTURA CHECK LIST PDF
# ══════════════════════════════════════════════════════════════════════════════
def leggi_checklist(codice_ui: str) -> str:
    """Estrae il testo dal check list PDF della UI (gestisce varianti del nome file)."""
    cartella = DATI_DIR / f"UI_{codice_ui}"
    # Prova le varianti di nome più comuni
    candidati = [
        cartella / f"UI_{codice_ui}_check list.pdf",
        cartella / f"UI_{codice_ui}_checklist.pdf",
        cartella / f"UI_{codice_ui}_check_list.pdf",
    ]
    pdf = next((p for p in candidati if p.exists()), None)
    if pdf is None:
        print(f"  ⚠  check list non trovata in {cartella}")
        return ""
    doc = fitz.open(str(pdf))
    testo = "\n".join(page.get_text() for page in doc)
    doc.close()
    return testo

# ══════════════════════════════════════════════════════════════════════════════
# CLAUDE API
# ══════════════════════════════════════════════════════════════════════════════
def chiedi_claude(prompt: str, max_tokens: int = 2000) -> str:
    resp = requests.post(
        "https://api.anthropic.com/v1/messages",
        headers={
            "Content-Type"    : "application/json",
            "x-api-key"       : ANTHROPIC_KEY,
            "anthropic-version": "2023-06-01",
        },
        json={
            "model"     : MODEL_CLAUDE,
            "max_tokens": max_tokens,
            "messages"  : [{"role": "user", "content": prompt}],
        },
        timeout=60,
    )
    resp.raise_for_status()
    return resp.json()["content"][0]["text"].strip()


def genera_voci_sezione_13(c: dict, checklist: str) -> list[str]:
    """
    Chiede a Claude di generare la lista puntata dei lavori per la sezione 1.3
    del PSC, basandosi sul contenuto della check list ITEA.
    Restituisce una lista di stringhe (una per voce).
    """
    ui   = c["codice_ui"]
    civ  = c["civico"]
    sc   = c["scala"]
    int_ = c["interno"]
    pn   = piano_str(c["piano"])
    sup  = c["sup"]
    tipo = c["tipo_intervento"]

    prompt = textwrap.dedent(f"""
        Sei un Coordinatore per la Sicurezza in fase di Esecuzione (CSE) e stai
        redigendo la sezione 1.3 «DESCRIZIONE DEGLI INTERVENTI DA REALIZZARE»
        di un PSC (Piano di Sicurezza e Coordinamento) ai sensi del D.Lgs. 81/2008.

        Dati del cantiere:
        - Unità immobiliare: UI_{ui}
        - Indirizzo: Via Doss Trento {civ}, scala {sc}, interno {int_}, piano {pn}
        - Superficie: {sup} m²
        - Tipo di intervento ITEA: {tipo}

        Di seguito il contenuto della check list ITEA per questa unità:
        ──────────────────────────────────────────────────────────────────
        {checklist[:3500]}
        ──────────────────────────────────────────────────────────────────

        Genera ESATTAMENTE 15 voci che descrivano le lavorazioni previste,
        in italiano tecnico-professionale tipico di un PSC.
        Ogni voce deve:
        • Iniziare direttamente con la descrizione (NO numeri, NO bullet iniziale)
        • Essere concisa (max 200 caratteri per voce)
        • Iniziare con il tipo di intervento/componente (es. "sostituzione dei serramenti...",
          "rifacimento parziale dell'impianto...", etc.)
        • Terminare con punto e virgola, tranne l'ultima che termina con punto.

        Se la check list ha meno di 15 interventi significativi, aggrega quelli
        affini in un'unica voce. Se ne ha di più, usa i principali.

        Rispondi SOLO con le 15 voci, una per riga, senza intestazioni né commenti.
    """).strip()

    risposta = chiedi_claude(prompt)

    # Parsing: rimuovi eventuali bullet/numeri/simboli iniziali
    voci = []
    for line in risposta.splitlines():
        line = line.strip()
        if not line:
            continue
        # Rimuovi bullet o numerazione iniziali
        line = re.sub(r'^[\d]+[.)]\s*', '', line)
        line = re.sub(r'^[•\-–*]\s*', '', line)
        if line:
            voci.append(line)

    return voci

# ══════════════════════════════════════════════════════════════════════════════
# SOSTITUZIONE TESTO IN DOCX (gestisce run spezzati)
# ══════════════════════════════════════════════════════════════════════════════
def _replace_in_para(para, old: str, new: str) -> bool:
    """
    Sostituisce old→new in un paragrafo python-docx.
    Prima tenta run per run; se fallisce (testo spezzato su più run)
    ricostruisce il testo completo nel primo run.
    Restituisce True se ha effettuato almeno una sostituzione.
    """
    if old not in para.text:
        return False

    # Tentativo 1: sostituzione in singolo run
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True

    # Tentativo 2: ricostruzione multi-run
    new_text = para.text.replace(old, new)
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""
    return True


def sostituisci_ovunque(doc, old: str, new: str) -> int:
    """Sostituisce old→new in tutto il documento: corpo, tabelle, header, footer."""
    if old == new:
        return 0
    count = 0

    def _walk_paras(paras):
        nonlocal count
        for p in paras:
            if _replace_in_para(p, old, new):
                count += 1

    def _walk_tables(tables):
        for tbl in tables:
            for row in tbl.rows:
                for cell in row.cells:
                    _walk_paras(cell.paragraphs)
                    _walk_tables(cell.tables)

    _walk_paras(doc.paragraphs)
    _walk_tables(doc.tables)

    for section in doc.sections:
        for hf in [section.header, section.footer,
                   section.first_page_header, section.first_page_footer]:
            if hf is not None:
                try:
                    _walk_paras(hf.paragraphs)
                    _walk_tables(hf.tables)
                except Exception:
                    pass

    return count


# ══════════════════════════════════════════════════════════════════════════════
# AGGIORNAMENTO SEZIONE 1.3
# ══════════════════════════════════════════════════════════════════════════════
def trova_bullet_sezione_13(doc) -> tuple[int, int]:
    """
    Restituisce (start_idx, end_idx) degli indici in doc.paragraphs
    che corrispondono alle voci bullet di sezione 1.3.
    start_idx = primo paragrafo bullet
    end_idx   = indice del primo paragrafo DOPO i bullet
    """
    paras = doc.paragraphs
    in_13 = False
    start = None
    end   = None

    for i, p in enumerate(paras):
        txt = p.text.strip()

        # Inizio sezione 1.3: cerchiamo "DESCRIZIONE DEGLI INTERVENTI"
        if "DESCRIZIONE DEGLI INTERVENTI" in txt.upper() and not in_13:
            in_13 = True
            continue

        if in_13:
            # Inizio voci: paragrafo che inizia con una lavorazione
            # (di solito segue un paragrafo intro "Le lavorazioni previste...")
            if start is None:
                # Il secondo paragrafo significativo dopo l'header 1.3 è l'intro;
                # il terzo in poi sono le voci bullet.
                # Identifichiamo le voci bullet come paragrafi brevi (< 300 car)
                # che NON iniziano con "Il " / "Le " / "L'" / "1." / "2."
                is_bullet = (
                    txt
                    and not txt.startswith("L'")
                    and not txt.startswith("Le ")
                    and not txt.startswith("Il ")
                    and not txt.startswith("1.")
                    and "FINALITÀ" not in txt.upper()
                    and len(txt) < 350
                )
                if is_bullet and ("sostituzione" in txt.lower()
                                  or "rifacimento" in txt.lower()
                                  or "tinteggiatura" in txt.lower()
                                  or "verifica" in txt.lower()
                                  or "interventi" in txt.lower()
                                  or "pulizia" in txt.lower()
                                  or "rimozione" in txt.lower()
                                  or "posa " in txt.lower()
                                  or "bonifica" in txt.lower()
                                  or "regolazione" in txt.lower()):
                    start = i

            else:
                # Fine voci: intestazione della sezione successiva
                if ("1.4" in txt or "ORGANIZZAZIONE" in txt.upper()
                        or "FASI LAVORATIVE" in txt.upper()):
                    end = i
                    break

    return start, end


def aggiorna_sezione_13(doc, nuove_voci: list[str]) -> None:
    """Sostituisce le voci bullet della sezione 1.3 con nuove_voci."""
    start, end = trova_bullet_sezione_13(doc)

    if start is None:
        print("  ⚠  Sezione 1.3 non trovata: aggiornamento saltato")
        return

    paras = doc.paragraphs
    bullet_paras = paras[start: end] if end else paras[start: start + 20]

    n_tmpl = len(bullet_paras)
    n_nuov = len(nuove_voci)
    print(f"     Sezione 1.3: {n_tmpl} voci template → {n_nuov} voci generate")

    for i, p in enumerate(bullet_paras):
        if i < n_nuov:
            if p.runs:
                p.runs[0].text = nuove_voci[i]
                for r in p.runs[1:]:
                    r.text = ""
            # Se non ci sono run, la voce resta com'è
        else:
            # Svuota le voci in eccesso
            for r in p.runs:
                r.text = ""


# ══════════════════════════════════════════════════════════════════════════════
# GENERAZIONE DI UN SINGOLO PSC
# ══════════════════════════════════════════════════════════════════════════════
def genera_psc(c: dict) -> Path:
    ui  = c["codice_ui"]
    civ = c["civico"]
    sc  = c["scala"]
    int_= c["interno"]
    pn  = piano_str(c["piano"])
    ped = c["p_ed"]
    sup = c["sup"]
    ac  = c["anno_costr"]
    gg  = c["durata_gg"]
    mesi= gg_to_mesi(gg)
    imp = fmt_euro(c["importo"])
    on  = fmt_euro(c["oneri"])
    di  = c["data_inizio"]

    print(f"\n{'═'*62}")
    print(f"  PSC  UI_{ui}  —  Via Doss Trento {civ}  sc.{sc} int.{int_} p.{pn}")
    print(f"{'═'*62}")

    out_path = OUTPUT_DIR / f"PSC_UI_{ui}.docx"
    shutil.copy(TEMPLATE_DOCX, out_path)
    doc = Document(str(out_path))

    # ── Sostituzioni (ordine: stringhe più lunghe/specifiche per prime) ──────
    sostituzioni = [

        # ── Codice UI ─────────────────────────────────────────────────────
        (f"UI_43103914 – Comune di Trento – Via Doss Trento 45/1 – scala A – interno 23 – piano 5°",
         f"UI_{ui} – Comune di Trento – Via Doss Trento {civ} – scala {sc} – interno {int_} – piano {pn}"),

        (f"Unità immobiliare UI_43103914",
         f"Unità immobiliare UI_{ui}"),

        (f"unità immobiliare UI_43103914",
         f"unità immobiliare UI_{ui}"),

        (f"manutenzione ordinaria e straordinaria – UI_43103914",
         f"manutenzione ordinaria e straordinaria – UI_{ui}"),

        ("AnalisiRischi_UI_43103914",    f"AnalisiRischi_UI_{ui}"),
        ("Cronoprogramma_UI_43103914",   f"Cronoprogramma_UI_{ui}"),
        ("OneriSicurezza_UI_43103914",   f"OneriSicurezza_UI_{ui}"),

        ("UI_43103914",  f"UI_{ui}"),
        ("43103914",     ui),

        # ── Indirizzo completo ─────────────────────────────────────────────
        (f"Via Doss Trento 45/1 – scala A – interno 23 – piano 5°",
         f"Via Doss Trento {civ} – scala {sc} – interno {int_} – piano {pn}"),

        (f"Via Doss Trento 45/1, scala A, interno 23, piano 5°",
         f"Via Doss Trento {civ}, scala {sc}, interno {int_}, piano {pn}"),

        (f"Lavori di manutenzione ordinaria e straordinaria – UI_43103914",
         f"Lavori di manutenzione ordinaria e straordinaria – UI_{ui}"),

        # Notifica preliminare T16[1,1]
        (f"Via Doss Trento 45/1 – scala A – interno 23 – piano 5°\n38122 TRENTO – C.C. Trento – p.ed. 5620",
         f"Via Doss Trento {civ} – scala {sc} – interno {int_} – piano {pn}\n38122 TRENTO – C.C. Trento – p.ed. {ped}"),

        # Sottostringa con dashes
        (f"– scala A – interno 23 – piano 5°",
         f"– scala {sc} – interno {int_} – piano {pn}"),

        (f"scala A, interno 23, piano 5°",
         f"scala {sc}, interno {int_}, piano {pn}"),

        (f"Via Doss Trento 45/1",
         f"Via Doss Trento {civ}"),
        (f"via Doss Trento 45/1",          # lowercase in body text
         f"via Doss Trento {civ}"),
        ("45/1",  civ),                     # cella tabella civico (valore isolato)

        # ── Componenti indirizzo singoli ───────────────────────────────────
        ("interno 23",   f"interno {int_}"),
        # Ordine: dal più specifico (con contesto) al più generico (valore isolato)
        # "al piano 5°" deve precedere "piano 5°" che deve precedere "5°" isolato
        ("al piano 5°",  f"al piano {pn}"),
        ("piano 5°",     f"piano {pn}"),
        # FIX: la cella tabella T1 "Piano" contiene solo "5°" senza la parola "piano"
        # Deve venire DOPO le sostituzioni più lunghe per non lasciare residui
        ("5°",           pn),
        ("p.ed. 5620",   f"p.ed. {ped}"),

        # Tabella T1 — valori cella singola
        ("958",   c["cod_ed"]),    # codice edificio
        ("5620",  ped),             # P.ED.
        ("99,81 m²", f"{sup} m²"),
        ("99,81",     sup),

        # ── Anno costruzione ───────────────────────────────────────────────
        ("anno di costruzione 1983",  f"anno di costruzione {ac}"),
        ("costruzione 1983",          f"costruzione {ac}"),
        ("edificio (1983)",           f"edificio ({ac})"),
        ("1983",                      ac),

        # ── Durata ────────────────────────────────────────────────────────
        (f"4 mesi naturali e consecutivi",
         f"{mesi} naturali e consecutivi"),

        (f"4 mesi (circa 120 giorni solari)",
         f"{mesi} (circa {gg} giorni solari)"),

        ("120 giorni solari", f"{gg} giorni solari"),
        ("120 giorni",        f"{gg} giorni"),
        ("circa 120",         f"circa {gg}"),
        ("4 mesi",            mesi),

        # ── Data inizio ────────────────────────────────────────────────────
        ("8 giugno 2026",  di),

        # ── Importo e oneri ────────────────────────────────────────────────
        ("€ 17.534,71",            f"€ {imp}"),
        ("17.534,71",              imp),
        ("€ 3.109,31",             f"€ {on}"),
        ("3.109,31",               on),
        ("tremilacentonove/31",    c["oneri_lettere"]),
    ]

    # Amianto: modifica la frase nel par. 3.1.3
    if c["amianto"]:
        sostituzioni += [
            (
                "non risulta segnalata la presenza di amianto nell'unità immobiliare",
                "è segnalata la presenza di amianto nell'unità immobiliare; "
                "la bonifica dovrà essere eseguita prima dell'inizio dei lavori edili "
                "da ditta specializzata e autorizzata ai sensi del D.Lgs. 81/2008"
            ),
        ]

    # Applica tutte le sostituzioni
    for old, new in sostituzioni:
        n = sostituisci_ovunque(doc, old, new)
        if n:
            print(f"  ✓  «{old[:50]}»  →  «{new[:50]}»  [{n}×]")

    # ── Fix diretto cella tabella T1 riga "Piano" (fallback robusto) ─────────
    # La sostituzione "5°" → pn può fallire se il valore è spezzato su run
    # diversi. Come backup troviamo la riga della tabella dove cells[0]=="Piano"
    # e impostiamo direttamente il testo di cells[1].
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = row.cells
            if len(cells) >= 2 and cells[0].text.strip() == "Piano":
                actual = cells[1].text.strip()
                if actual != pn:
                    # Imposta il valore corretto
                    for para in cells[1].paragraphs:
                        for run in para.runs:
                            run.text = ""
                    cells[1].paragraphs[0].runs[0].text = pn if cells[1].paragraphs[0].runs else pn
                    if not cells[1].paragraphs[0].runs:
                        cells[1].paragraphs[0].add_run(pn)
                    print(f"  ✓  Cella tabella Piano: «{actual}» → «{pn}» (fix diretto)")

    # ── Sezione 1.3: genera da check list ─────────────────────────────────
    print(f"  Lettura check list UI_{ui}…")
    checklist = leggi_checklist(ui)

    if checklist:
        print(f"  Chiamata Claude per sezione 1.3…")
        try:
            voci = genera_voci_sezione_13(c, checklist)
            aggiorna_sezione_13(doc, voci)
        except Exception as exc:
            print(f"  ⚠  Errore Claude sezione 1.3: {exc}")
            print(     "     La sezione 1.3 manterrà il testo template")
    else:
        print("  ⚠  Check list assente — sezione 1.3 lasciata dal template")

    # ── Salva ──────────────────────────────────────────────────────────────
    doc.save(str(out_path))
    print(f"  ✓  Salvato: {out_path.name}")
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print()
    print("╔══════════════════════════════════════════════════════════════╗")
    print("║  Generazione batch PSC — ITEA Via Doss Trento, Trento       ║")
    print("╚══════════════════════════════════════════════════════════════╝")
    print(f"  Template : {TEMPLATE_DOCX.name}")
    print(f"  Output   : {OUTPUT_DIR}")
    print()

    if not TEMPLATE_DOCX.exists():
        sys.exit(f"ERRORE: template non trovato:\n  {TEMPLATE_DOCX}")

    generati, errori = [], []

    for cantiere in CANTIERI:
        try:
            out = genera_psc(cantiere)
            generati.append(out.name)
        except Exception as exc:
            print(f"  ✗  ERRORE per UI_{cantiere['codice_ui']}: {exc}")
            errori.append(cantiere["codice_ui"])

    print()
    print("══════════════════════════════════════════════════════════════")
    print(f"  Completato: {len(generati)} / {len(CANTIERI)} PSC generati")
    if errori:
        print(f"  Errori UI: {', '.join(errori)}")
    print(f"  Output  → {OUTPUT_DIR}")
    print("══════════════════════════════════════════════════════════════")
